import os
import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from PDF, DOCX, or TXT files.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text from the file
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
        Exception: For other processing errors
    """
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt':
            return extract_text_from_txt(file_path)
        elif file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
            
    except Exception as e:
        raise Exception(f"Error extracting text from {file_path}: {str(e)}")


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber (more reliable) with PyPDF2 fallback."""
    text = ""
    
    try:
        # Method 1: Using pdfplumber (more accurate for complex layouts)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            return text.strip()
            
    except Exception as e:
        print(f"pdfplumber failed: {e}, trying PyPDF2...")
    
    try:
        # Method 2: Fallback to PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return text.strip()
        
    except Exception as e:
        raise Exception(f"Both PDF extraction methods failed: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(file_path)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
        
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


# Alternative function with more detailed error handling
def extract_text_with_metadata(file_path: str) -> dict:
    """
    Extract text with metadata about the extraction process.
    
    Returns:
        dict: {
            'text': str,
            'file_type': str,
            'file_size': int,
            'status': str,
            'pages': int (for PDF),
            'error': str (if any)
        }
    """
    result = {
        'text': '',
        'file_type': '',
        'file_size': 0,
        'status': 'success',
        'pages': 0,
        'error': None
    }
    
    try:
        if not os.path.exists(file_path):
            result['status'] = 'error'
            result['error'] = f"File not found: {file_path}"
            return result
        
        result['file_size'] = os.path.getsize(file_path)
        result['file_type'] = os.path.splitext(file_path)[1].lower()
        
        # Extract text
        text = extract_text_from_file(file_path)
        result['text'] = text
        
        # Get page count for PDF
        if result['file_type'] == '.pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    result['pages'] = len(pdf.pages)
            except:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        result['pages'] = len(pdf_reader.pages)
                except:
                    result['pages'] = 0
        
        return result
        
    except Exception as e:
        result['status'] = 'error'
        result['error'] = str(e)
        return result


